from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def add_para(text, bold=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    return p

def add_table_with_style(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Microsoft YaHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Microsoft YaHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ==================== COVER PAGE ====================
doc.add_paragraph()
doc.add_paragraph()
add_para('Tokenhub — 企业AI治理智能平台', bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('商业计划书', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0, 102, 204))
doc.add_paragraph()
add_para('融资轮次：天使轮  |  融资需求：¥500万  |  出让股权：待定', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_para('2026年6月', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ==================== TABLE OF CONTENTS (placeholder) ====================
add_heading_styled('目录', 1)
toc_items = [
    '一、项目概述',
    '二、市场机会与规模',
    '三、行业痛点',
    '四、产品方案',
    '五、商业模式',
    '六、竞争分析',
    '七、技术亮点',
    '八、发展路线图',
    '九、运营数据与里程碑',
    '十、团队介绍',
    '十一、融资需求与资金用途',
    '十二、风险与应对',
]
for item in toc_items:
    add_para(item, size=12)
doc.add_page_break()

# ==================== 1. PROJECT OVERVIEW ====================
add_heading_styled('一、项目概述', 1)

add_para('1.1 一句话定位', bold=True, size=12)
add_para('Tokenhub 是一款面向企业客户的 AI 治理与 Token 安全运营中台，帮助企业在安全、合规、可控的前提下高效使用大模型，同时降低30%以上的 AI 调用成本。')

add_para('1.2 产品形态', bold=True, size=12)
add_para('SaaS + 私有化部署双模式。目前已接入8家主流模型厂商（DeepSeek、通义千问、豆包、智谱GLM、Kimi、文心一言、混元等），覆盖9个主力模型，提供统一 API 网关、智能路由、内容安全过滤、预算管控、合规审计、Agent 安全管理、智能体市场、知识库安全治理等全栈能力。')

add_para('1.3 当前阶段', bold=True, size=12)
add_para('MVP（最小可行产品）已完成开发并通过测试验证：')
add_para('— 后端 15 个核心模块、70+ API 端点，采用 Go 语言 + Gin 框架')
add_para('— 前端 14 个功能页面，采用 React/Next.js + Ant Design')
add_para('— Go 单元测试 30+、前端单元测试 40+、E2E 测试 6+，全部通过')
add_para('— 当前可演示（Dashboard、安全治理、Agent 市场、知识库管理等）')

add_para('1.4 核心价值主张', bold=True, size=12)
add_para('帮企业省钱：智能路由 + 预算熔断，降低 AI 调用成本 30-50%')
add_para('帮企业避险：四层安全引擎（敏感词 + PII + 注入检测 + DLP），杜绝数据泄露')
add_para('帮企业合规：使用台账 + 安全审计 + 算法备案，一站式满足监管要求')
add_para('帮企业提效：统一接入 8 家模型厂商，告别多 Key、多平台管理混乱')

doc.add_page_break()

# ==================== 2. MARKET OPPORTUNITY ====================
add_heading_styled('二、市场机会与规模', 1)

add_para('2.1 市场驱动力', bold=True, size=12)
add_para('当前全球 AI 产业正处于爆发期，企业对大模型的采纳速度前所未有，但随之而来的治理、安全和成本管理需求也急剧增长。以下四大趋势共同推动 AI 治理平台成为刚需：')
add_para('1) 企业 AI 采纳率激增：据中国信通院数据，2025年中国企业AI采纳率已超过45%，大型企业更是高达78%。模型数量从2023年的平均1.5个增长到2025年的4.3个/企业。')
add_para('2) Token 消耗爆发式增长：单个企业月度 Token 消耗从2023年的平均50万增长至2025年的800万+，AI 支出成为企业IT预算中增速最快的项目。')
add_para('3) 安全事件频发：三星、亚马逊、摩根大通等企业已发生多起因 AI 使用导致的数据泄露事件。Gartner 预测，到2026年，50%的企业将因 AI 安全事件遭受重大财务或声誉损失。')
add_para('4) 监管持续收紧：中国《生成式人工智能服务管理暂行办法》已正式实施，算法备案、数据安全评估、内容安全审核成为企业使用 AI 的硬性门槛。')

add_para('2.2 市场规模', bold=True, size=12)
add_para('以下数据来自权威第三方研究机构，展示 AI 治理及相关市场的规模与增长趋势：')

add_table_with_style(
    ['市场领域', '2024年规模', '2028年预测规模', 'CAGR', '数据来源'],
    [
        ['全球 AI 治理市场', '$1.8B', '$6.2B', '28.2%', 'MarketsandMarkets 2025'],
        ['中国 AI 安全市场', '¥48亿', '¥186亿', '31.1%', 'IDC China 2025'],
        ['API 管理/网关市场', '$5.4B', '$17.3B', '26.1%', 'Gartner 2025'],
        ['企业 AI 平台市场(中国)', '¥132亿', '¥508亿', '30.9%', '艾瑞咨询 2025'],
        ['全球 AI 网络安全市场', '$25.5B', '$50.8B', '14.8%', 'MarketsandMarkets 2026'],
        ['AI Agent 市场', '$5.1B', '$47.1B', '55.8%', 'MarketsandMarkets 2025'],
    ],
    [3.5, 2.5, 2.8, 2.2, 3.0]
)

add_para('')
add_para('2.3 目标市场测算（TAM/SAM/SOM）', bold=True, size=12)
add_table_with_style(
    ['维度', '定义', '规模', '测算依据'],
    [
        ['TAM（总可触达市场）', '中国 AI 治理与安全市场', '¥186亿（2028E）', 'IDC China × 企业AI平台规模'],
        ['SAM（可服务市场）', '中大型企业 AI 治理平台', '¥37亿', 'TAM × 20% (企业级治理细分)'],
        ['SOM（可获得市场）', '天使轮目标客户', '¥3,700万', '前3年目标：签约100家企业客户'],
    ],
    [3, 4, 3, 5]
)

doc.add_page_break()

# ==================== 3. PAIN POINTS ====================
add_heading_styled('三、行业痛点', 1)

add_para('基于 TaaS（Token-as-a-Service）安全企业实践洞察，当前企业落地 AI 面临六大核心困境：', size=11)

add_table_with_style(
    ['#', '痛点', '具体表现', '影响'],
    [
        ['1', '模型碎片化', '多厂商、多模型、多 API Key 分散管理，缺乏统一接入入口', '管理成本高，安全风险点多'],
        ['2', '成本黑洞', 'Token 消耗不透明，缺乏分部门/项目核算，预算频繁超支', 'AI 支出 ROI 难以衡量'],
        ['3', '安全风险', '提示词注入攻击、敏感数据泄露、API Key 盗用、越权调用', '品牌声誉受损，可能面临法律诉讼'],
        ['4', 'Agent 失控', 'Agent 自主循环消耗 Token、工具链投毒、执行越权操作', '资源恶性消耗，业务流程失控'],
        ['5', '合规压力', '等保三级、算法备案、数据分类分级、生成式AI管理办法', '违规使用可能被责令整改或下架'],
        ['6', '计量纠纷', '多模型/多供应商/多模态下计费口径不一致，对账争议频发', '供应商关系紧张，结算争议增加运营成本'],
    ],
    [0.5, 2, 6, 6]
)

add_para('')
add_para('这些痛点在当前市场上缺乏一站式的解决方案。现有产品要么只做API网关（缺安全合规）、要么只做内容安全（缺预算管控）、要么只做成本优化（缺Agent治理）。Tokenhub 是市场上首个覆盖"供给侧→平台侧→消费侧→智能体侧"全链路的AI治理平台。')

doc.add_page_break()

# ==================== 4. PRODUCT ====================
add_heading_styled('四、产品方案', 1)

add_para('4.1 产品架构', bold=True, size=12)
add_para('Tokenhub 采用三层架构设计，覆盖 TaaS 全生态链：')

add_para('第一层：供给侧 — 模型聚合与交付', bold=True)
add_para('— 8家厂商 9个模型统一适配（DeepSeek V3/R1、通义千问 Max/Plus、豆包 Pro 256K、GLM-5/5-Flash/4-Plus、Kimi、文心4.5、混元 Pro）')
add_para('— 统一 OpenAI 兼容 API 接口，零代码切换模型')
add_para('— 智能路由引擎（代价优先/质量优先/均衡三种策略，自动 fallback）')

add_para('第二层：平台侧 — 安全运营中枢', bold=True)
add_para('— Token 生命周期管理（发行→激活→消费→冻结→销毁）')
add_para('— 可信计量账本（全量元数据 + 签名摘要，不可篡改）')
add_para('— 统一安全网关（身份认证 + DLP + 敏感词 + PII + 注入检测 + 限流）')
add_para('— 五级预算管控 + 三级熔断器 + Agent Guard 异常防护')

add_para('第三层：消费侧 — 企业内部治理', bold=True)
add_para('— 统一接入网关（强制所有业务通过平台调用 AI）')
add_para('— 场景登记 + 数据分级管控')
add_para('— AI FinOps（成本归集/预算阈值/模型选型优化/缓存压缩/异常复盘）')
add_para('— 知识库安全治理（上传即扫描，联动安全巡检 Agent）')

add_para('4.2 核心功能矩阵', bold=True, size=12)
add_table_with_style(
    ['模块', '功能', '对标 TaaS 标准', '状态'],
    [
        ['模型方舟', '8厂商9模型统一接入 + 智能路由 + 场景推荐', '供给侧可信交付', 'MVP 已完成'],
        ['安全网关', '敏感词14类 + PII识别脱敏 + 注入检测 + DLP', '消费侧安全基线', 'MVP 已完成'],
        ['预算管控', '5级预算(公司/部门/项目/用户/Key) + 3态熔断 + Agent Guard', 'AI FinOps 最佳实践', 'MVP 已完成'],
        ['安全巡检Agent', '注入/权限/合规/供应链4维自动扫描', '智能体侧治理', 'MVP 已完成'],
        ['智能体市场', '专家/Skills/Connectors/MCP 工具生态', 'TaaS 生态全景', 'MVP 已完成'],
        ['知识库管理', '上传即扫描，安全联动（4维度检测）', '供给侧可信交付', 'MVP 已完成'],
        ['合规报告', '使用台账/安全审计/算法备案 三大报告', '评估侧合规', 'MVP 已完成'],
        ['沙箱审核', 'Agent 输出人工审核 + 自动化规则', '智能体侧治理', 'MVP 已完成'],
        ['Token Broker', '短期动态凭证 + 即时撤销 + 细粒度授权', '平台侧凭证管理', 'V1.0 规划'],
        ['计量账本', '不可篡改交易数据底座 + 全量元数据签名', '平台侧可信计量', 'V1.0 规划'],
        ['交易平台', '权益流转 + 合作伙伴管理 + 争议处置', '延伸实践', '长期规划'],
    ],
    [2, 5, 4, 2.5]
)

doc.add_page_break()

# ==================== 5. BUSINESS MODEL ====================
add_heading_styled('五、商业模式', 1)

add_para('5.1 定价体系', bold=True, size=12)
add_table_with_style(
    ['版本', '价格(年)', '核心功能', '目标客户'],
    [
        ['社区版', '免费', '10个API Key + 基础安全 + 3个模型', '个人开发者/初创'],
        ['专业版', '¥9.8万', '无限Key + 高级安全 + FinOps + 知识库', '中小企业(50-500人)'],
        ['企业版', '¥29.8万', '私有部署 + SSO + 定制合规 + Agent市场', '政企大客户(500人+)'],
        ['旗舰版', '¥68万+', '专属Agent + 交易平台 + 专属运维', '生态运营商/超大型企业'],
    ],
    [2.5, 2.5, 6, 4]
)

add_para('')
add_para('5.2 附加收入来源', bold=True, size=12)
add_para('— Agent 市场佣金：专家被订阅时抽取 15-20% 佣金')
add_para('— 安全评估服务：对企业AI系统进行渗透测试和安全审计（¥5-15万/次）')
add_para('— 定制 Agent 开发：为企业定制专属安全巡检 Agent 和业务 Agent')
add_para('— API 开放平台：未来向第三方开发者开放 API 接口，按调用量收费')

add_para('5.3 盈利预测（3年）', bold=True, size=12)
add_table_with_style(
    ['年份', '客户数', 'ARR(万元)', '毛利率', '净利率'],
    [
        ['Y1 (2026-2027)', '30', '350', '65%', '-30%（投入期）'],
        ['Y2 (2027-2028)', '80', '1,500', '72%', '8%'],
        ['Y3 (2028-2029)', '200', '5,000', '78%', '22%'],
    ],
    [3, 2, 3, 2, 3]
)

doc.add_page_break()

# ==================== 6. COMPETITION ====================
add_heading_styled('六、竞争分析', 1)

add_para('6.1 主要竞品对比', bold=True, size=12)
add_para('以下选取三家直接/间接竞品进行全方位对比：')

add_table_with_style(
    ['维度', 'Tokenhub', '硅基流动\n(SiliconFlow)', '天翼云 Tokenhub\n(息壤)', '阿里云\n(百炼/PAI)'],
    [
        ['定位', '企业AI治理 + 安全中台', '模型API聚合 + 推理加速', '算力调度 + Token运营', '全栈AI中台'],
        ['模型接入', '8厂商9模型\n(多厂商聚合)', '自研 + 开源模型\n(30+模型)', '天翼自研模型为主\n(星辰系列)', '阿里自研模型为主\n(通义系列)'],
        ['安全检测', '4层引擎\n(敏感词/PII/注入/DLP)', '基础API Key管理', '运营商级安全\n(网络/链路层)', '内容安全\n(通义安全)'],
        ['预算管控', '5级预算+3态熔断\n+Agent Guard', '基础计费/限额', '算力配额管理', '资源组/配额管理'],
        ['Agent治理', '独立身份+安全巡检\n+沙箱审核', '无', '无', 'Agent开发平台'],
        ['知识库安全', '上传即扫描\n4维度检测', '无', '无', '无'],
        ['合规报告', '台账/审计/备案\n三大报告', '基础使用统计', '基础用量报表', '基础日志/审计'],
        ['TaaS生态', '全链路(供/平/消/智)', '仅供给侧', '供给侧+平台侧', '仅平台侧'],
        ['部署方式', 'SaaS + 私有化', 'SaaS API', '私有化为主', '公有云 + 混合云'],
        ['定价', '免费版+¥9.8万起', '按Token计费\n(有免费额度)', '按算力计费', '按Token/资源组计费'],
        ['目标客户', '中大型企业\n(金融/政务/制造)', 'AI开发者\n/中小应用', '政企客户\n(电信行业)', '阿里云存量客户\n/大型企业'],
        ['核心壁垒', '全链路AI治理\n唯一平台', '推理加速性能\n/模型丰富', '运营商渠道+\n算力基础设施', '云生态绑定\n/品牌+渠道'],
    ],
    [2, 2.8, 3, 3, 3]
)

add_para('')
add_para('6.2 差异化优势总结', bold=True, size=12)
add_para('1) 唯一全链路覆盖：市场上仅有的同时覆盖"供给侧→平台侧→消费侧→智能体侧"的治理平台，单一竞品通常只覆盖1-2个环节。')
add_para('2) 安全深度领先：Aho-Corasick 14类敏感词检测 + PII 脱敏 + 注入防护 + DLP 四层引擎，远超行业水平。')
add_para('3) Agent 安全巡检独创：对 Agent/Skill/MCP 进行注入风险、权限审计、合规检查、供应链漏洞四维度自动扫描，行业首创。')
add_para('4) 厂商中立：不绑定任何模型厂商，客户可自由选择最优模型组合，避免供应商锁定。')
add_para('5) 知识库安全联动：上传即检测，将企业知识库纳入安全治理体系，解决 RAG 场景安全盲区。')

doc.add_page_break()

# ==================== 7. TECHNICAL HIGHLIGHTS ====================
add_heading_styled('七、技术亮点', 1)

add_para('7.1 技术架构', bold=True, size=12)
add_table_with_style(
    ['层次', '技术选型', '说明'],
    [
        ['后端语言/框架', 'Go 1.22 + Gin v1.10', '高并发网关，支持万级 QPS'],
        ['前端框架', 'React 19 + Next.js 15 + Ant Design 5', '现代化管理后台，SSR/CSR 混合'],
        ['数据库', 'PostgreSQL + ClickHouse + Redis', 'OLTP + OLAP + 缓存 三层存储'],
        ['安全引擎', 'Aho-Corasick 自动机', '14类敏感词毫秒级匹配'],
        ['测试', 'Go test + Vitest + Playwright', '30+单元测试 + 40+前端测试 + 6+E2E'],
    ],
    [3, 5, 7]
)

add_para('')
add_para('7.2 核心技术能力', bold=True, size=12)
add_para('— 多模型适配器架构：统一接口抽象，新模型接入 < 1天（对接 OpenAI 兼容 API）')
add_para('— Aho-Corasick 敏感词引擎：14类敏感词实时检测（涉政/色情/暴力/赌博/违禁品/诈骗/暴恐/未成年人保护等），微秒级响应')
add_para('— Agent 安全巡检系统：配置变更事件驱动 + 10秒防抖 + 4技能并行扫描')
add_para('— 预算熔断器：Redis Lua 脚本原子操作，三级状态（normal→throttled→blocked），毫秒级判断')
add_para('— 智能路由引擎：基于8种业务场景的关键词分类推荐，支持 primary→fallback 降级链')

doc.add_page_break()

# ==================== 8. ROADMAP ====================
add_heading_styled('八、发展路线图', 1)

add_table_with_style(
    ['阶段', '时间', '重点任务', '里程碑'],
    [
        ['MVP 验证期', '2026 Q2 (当前)', '— 产品可演示\n— 种子用户试用\n— 收集反馈迭代', '10家试用客户签约\nNPS > 40'],
        ['V1.0 商用版', '2026 Q3-Q4', '— Token Broker 托管\n— 不可篡改计量账本\n— 凭证托管 + SLA 体系\n— RAG 知识库安全增强', '30家付费客户\nARR ¥200万'],
        ['生态扩展期', '2027 Q1-Q2', '— 权益交易平台\n— 合作伙伴额度管理\n— 第三方评估接口\n— API 开放平台', '80家客户\nARR ¥1,500万\nAgent市场100+专家'],
        ['行业深耕期', '2027 Q3-Q4', '— 金融合规版（等保三级+信创）\n— 政务定制版（XC适配）\n— 制造业方案\n— 海外多语言版', '200家客户\nARR ¥5,000万\n行业前3'],
        ['规模化增长', '2028+', '— 生态运营商模式\n— 行业ISV合作\n— ISO27001/SOC2 认证\n— 国际市场拓展', '500+客户\nARR ¥1.5亿+\nIPO 准备'],
    ],
    [2.5, 2.5, 5, 5]
)

doc.add_page_break()

# ==================== 9. MILESTONES ====================
add_heading_styled('九、运营数据与里程碑', 1)

add_para('9.1 产品完成度', bold=True, size=12)
add_table_with_style(
    ['指标', '数据', '行业对比'],
    [
        ['API 端点数', '70+', '同类产品平均 20-30'],
        ['接入模型数', '8厂商 9模型', '同类产品 2-5 模型'],
        ['安全检测层数', '4层（敏感词/PII/注入/DLP）', '同类产品 1-2 层'],
        ['代码规模', '~8,000行 Go + ~7,000行 TS', '—'],
        ['Go 单元测试', '30+（5包，100%通过）', '—'],
        ['前端组件测试', '40+（100%通过）', '—'],
        ['E2E 测试', '6+（100%通过）', '—'],
        ['TypeScript', '0 类型错误', '—'],
    ],
    [4, 5, 5]
)

add_para('')
add_para('9.2 KPI 目标', bold=True, size=12)
add_table_with_style(
    ['目标', '6个月', '12个月', '24个月', '36个月'],
    [
        ['付费客户', '10', '30', '80', '200'],
        ['ARR', '¥100万', '¥350万', '¥1,500万', '¥5,000万'],
        ['月活企业', '20', '50', '120', '300'],
        ['Agent 市场专家数', '15', '50', '100+', '300+'],
        ['NPS', '>40', '>45', '>50', '>55'],
        ['团队规模', '8人', '15人', '30人', '50人'],
    ],
    [2.5, 2, 2.5, 3, 3]
)

doc.add_page_break()

# ==================== 10. TEAM ====================
add_heading_styled('十、团队介绍', 1)

add_para('注：以下为团队简介模板，请自行填写具体信息。', size=10, color=(128, 128, 128))

team = [
    ('创始人 / CEO', '[姓名]', '[教育背景]', '[工作经历和核心成就]\n[在AI/企业服务/安全领域的经验]\n[创始动机与愿景]'),
    ('联合创始人 / CTO', '[姓名]', '[教育背景]', '[技术背景、主导的技术项目]\n[全栈能力、Go/React/安全领域经验]\n[技术架构设计能力]'),
    ('产品负责人', '[姓名]', '[教育背景]', '[产品背景，参与过的AI/企业服务产品]\n[TaaS行业标准参与经验]\n[对企业AI治理的理解]'),
    ('技术顾问', '[姓名]', '[教育背景]', '[学术或行业背景]\n[AI安全/企业服务领域的影响力]'),
]

for role, name, edu, desc in team:
    add_para(f'{role}：{name}', bold=True, size=12)
    add_para(f'教育背景：{edu}')
    add_para(f'核心经验：{desc}')
    add_para('')

doc.add_page_break()

# ==================== 11. FUNDING ====================
add_heading_styled('十一、融资需求与资金用途', 1)

add_para('11.1 融资方案', bold=True, size=12)
add_table_with_style(
    ['项目', '内容'],
    [
        ['融资轮次', '天使轮'],
        ['融资金额', '¥500万（人民币）'],
        ['出让股权', '待协商（建议 10-15%）'],
        ['投前估值', '¥3,300万 — ¥5,000万'],
        ['投资方式', '增资扩股'],
        ['资金到位时间', '2026年Q3'],
        ['预计 Runway', '18个月'],
    ],
    [4, 12]
)

add_para('')
add_para('11.2 资金用途', bold=True, size=12)
add_table_with_style(
    ['用途', '金额(万元)', '占比', '说明'],
    [
        ['产品研发', '200', '40%', 'Token Broker、计量账本、RAG安全模块、SLA体系建设\n扩充研发团队（+3名后端 + 2名前端）'],
        ['市场拓展', '150', '30%', '首批10家标杆客户落地\n行业会议/展会参与\n销售团队组建（+2名BD）'],
        ['安全合规认证', '75', '15%', '等保三级测评\nISO27001 信息安全管理体系认证\n算法备案'],
        ['运营储备', '75', '15%', '服务器/云服务费用（18个月）\n办公/行政/法务\n现金流安全垫'],
        ['合计', '500', '100%', ''],
    ],
    [3, 2.5, 2, 8]
)

add_para('')
add_para('11.3 投资回报预期', bold=True, size=12)
add_para('— 3年后 ARR 目标：¥5,000万，对应估值 ¥3.75-5亿（7.5-10x P/S）')
add_para('— 预计退出路径：A轮融资（2027-2028）、战略并购（大型云厂商/安全厂商）、或2029年启动 IPO 准备')
add_para('— 天使轮投资人预期回报：3年内 10-15x')

doc.add_page_break()

# ==================== 12. RISKS ====================
add_heading_styled('十二、风险与应对', 1)

add_table_with_style(
    ['风险类别', '风险描述', '应对策略'],
    [
        ['市场竞争', '大型云厂商（阿里云、华为云）推出类似产品', '差异化竞争：不绑定模型厂商、安全深度领先、全链路覆盖'],
        ['技术壁垒', '模型厂商自建安全能力，削弱平台价值', '向上游延伸：Agent 市场、知识库治理、合规报告等增值服务'],
        ['客户获取', '企业采购周期长，获客成本高', 'PLG 策略：社区版免费 → 专业版转化；标杆客户案例驱动'],
        ['安全合规变化', 'AI监管政策持续演进，合规要求提高', '模块化设计：合规模块可独立升级；政策预研团队'],
        ['资金风险', '研发和市场拓展超出预算', '严格控制 burn rate，18个月 runway；优先高毛利企业版客户'],
        ['人才风险', 'AI安全人才稀缺，招聘困难', '股权激励 + 远程灵活办公 + 技术品牌建设（开源/技术博客）'],
    ],
    [2.5, 5, 5]
)

# save
output_dir = r'C:\Users\chenjiawei\Desktop\Tokenhub\docs\bp'
os.makedirs(output_dir, exist_ok=True)
doc_path = os.path.join(output_dir, 'Tokenhub_投资BP_商业计划书.docx')
doc.save(doc_path)
print(f'Word saved: {doc_path}')
